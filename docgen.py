"""
docgen.py
Módulo encargado de toda la lógica relacionada con python-docx:
- BORRADO condicional de bloques según marcadores
- REEMPLAZO de marcadores {KEY}
- INSERCIÓN de hallazgos
- INSERCIÓN de tabla de muestreo
- Función pública: generate_report(template_bytes, user_data, hallazgos, muestreo_rows)
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
import re
import io

# ---------------------------
# CONSTANTES (copiadas/adaptadas de tu app)
# ---------------------------
CAAP_LOGICA_ESTADOS = [
    "No aplica",
    "No iniciado, tiene CNCA vigente",
    "No iniciado, tiene CNCA en curso",
    "No iniciado, tiene CNCA vencida / No solicitada)",
    "En curso",
    "Vigente",
    "Vencido"
]
ADA_DETALLE_ESTADOS = ["Seleccione...", "Prefactibilidad con Chi 0 para todos los permisos", "Prefactibilidad vigente", "Prefactibilidad vencida", "No solicitada"]
RENPRE_DETALLE_ESTADOS = ["Seleccione...", "No aplica", "Esta inscripto y renueva", "Esta inscripto y no renovó", "Aplica pero no esta inscripto"]
FRECUENCIA_MUESTREO_OPTIONS = ["Seleccione...", "Anual", "Semestral", "Trimestral", "Mensual", "Continua", "N/A"]

RESIDUOS_ESPECIALES_STATUS_GENERAL = [
    "Seleccione...", "Empresa exenta porque no genera", "No inscripta", "No Cumple con las DDJJ",
    "Cumple con las DDJJ"
]
CHE_DETALLE_ESTADOS = [
    "Seleccione...", "Obtuvo CHE", "No obtuvo CHE", "CHE en curso"
]
ACUMAR_DETALLE_ESTADOS = ["Seleccione...", "No aplica", "Vigente", "En curso", "Vencida", "No solicitada"]
SE_DETALLE_ESTADOS = ["Seleccione...", "No aplica", "Vigente", "En curso", "Vencida", "No solicitada"]

# MARCADORES_CONDICIONALES (adaptado)
MARCADORES_CONDICIONALES = {
    "HABILITACION_MUNICIPAL": {
        "cumple": {"start": "{INICIO_HAB_CUMPLE}", "end": "{FIN_HAB_CUMPLE}"},
        "no cumple": {"start": "{INICIO_HAB_NO_CUMPLE}", "end": "{FIN_HAB_NO_CUMPLE}"},
        "parcial": {"start": "{INICIO_HAB_PARCIAL}", "end": "{FIN_HAB_PARCIAL}"
    },
    "CNCA_STATUS": {
        "vigente": {"start": "{INICIO_CNCA_VIGENTE}", "end": "{FIN_CNCA_VIGENTE}"},
        "superada (esta en curso el CAA)": {"start": "{INICIO_CNCA_SUPERADA}", "end": "{FIN_CNCA_SUPERADA}"},
        "vencida": {"start": "{INICIO_CNCA_VENCIDA}", "end": "{FIN_CNCA_VENCIDA}"},
        "no solicitada": {"start": "{INICIO_CNCA_NO_SOLICITADA}", "end": "{FIN_CNCA_NO_SOLICITADA}"},
        "en curso": {"start": "{INICIO_CNCA_EN_CURSO}", "end": "{FIN_CNCA_EN_CURSO}"
    },
    "CAAP_STATUS": {
        "No iniciado, tiene CNCA vigente": {"start": "{INICIO_CAAP_NO_INICIADO_CNCA_VIGENTE}", "end": "{FIN_CAAP_NO_INICIADO_CNCA_VIGENTE}"},
        "No iniciado, tiene CNCA en curso": {"start": "{INICIO_CAAP_NO_INICIADO_CNCA_ENCURSO}", "end": "{FIN_CAAP_NO_INICIADO_CNCA_ENCURSO}"},
        "No iniciado, tiene CNCA vencida / No solicitada)": {"start": "{INICIO_CAAP_NO_INICIADO_CNCA_VENCIDA_NO_SOLICITADA}", "end": "{FIN_CAAP_NO_INICIADO_CNCA_VENCIDA_NO_SOLICITADA}"},
        "En curso": {"start": "{INICIO_CAAP_EN_CURSO}", "end": "{FIN_CAAP_EN_CURSO}"},
        "Vigente": {"start": "{INICIO_CAAP_VIGENTE}", "end": "{FIN_CAAP_VIGENTE}"},
        "Vencido": {"start": "{INICIO_CAAP_VENCIDO}", "end": "{FIN_CAAP_VENCIDO}"
    },
    "CAAF_STATUS": {
        "en_curso": {"start": "{INICIO_CAAF_EN CURSO}", "end": "{FIN_CAAF_EN_CURSO}"},
        "vigente": {"start": "{INICIO_CAAF_VIGENTE}", "end": "{FIN_CAAF_VIGENTE}"},
        "no_iniciado_caap_en_curso": {"start": "{INICIO_CAAF_NO_INICIADO_CAAP_EN CURSO}", "end": "{FIN_CAAF_NO_INICIADO_CAAP_EN_CURSO}"},
        "no_iniciado_caap_vencido": {"start": "{INICIO_CAAF_NO_INICIADO_CAAP_VENCIDO}", "end": "{FIN_CAAF_NO_INICIADO_CAAP_VENCIDO}"},
        "vencido": {"start": "{INICIO_CAAF_VENCIDO}", "end": "{FIN_CAAF_VENCIDO}"},
        "eliminar_todo": {"start": "{ELIMINAR_TODO_CAAF}", "end": "{FIN_ELIMINAR_TODO_CAAF}"
    },
    "RENOVACION_CAA_STATUS": {
        "En curso": {"start": "{INICIO_RENOVACION_CAA_ENCURSO}", "end": "{FIN_RENOVACION_CAA_ENCURSO}"
    },
    "LEGA_STATUS": {
        "vigente": {"start": "{INICIO_LEGA_VIGENTE}", "end": "{FIN_LEGA_VIGENTE}"},
        "en_curso": {"start": "{INICIO_LEGA_EN_CURSO}", "end": "{FIN_LEGA_EN_CURSO}"},
        "vencida": {"start": "{INICIO_LEGA_VENCIDA}", "end": "{FIN_LEGA_VENCIDA}"
    },
    "RESIDUOS_ESPECIALES_STATUS": {
        "Empresa exenta porque no genera": {"start": "{INICIO_EMPRESA_EXCENTA}", "end": "{FIN_EMPRESA_EXCENTA}"},
        "No inscripta": {"start": "{INICIO_RREE_NO_INSCRIPTA}", "end": "{FIN_RREE_NO_INSCRIPTA}"},
        "No Cumple con las DDJJ": {"start": "{INICIO_RREE_NOCUMPLE_DDJJ}", "end": "{FIN_RREE_NOCUMPLE_DDJJ}"},
        "Cumple con las DDJJ": {"start": "{INICIO_RREE_CUMPLE_DDJJ}", "end": "{FIN_RREE_CUMPLE_DDJJ}"
    },
    "CHE_STATUS": {
        "Obtuvo CHE": {"start": "{INICIO_RREE_OBTUVO_CHE}", "end": "{FIN_RREE_OBTUVO_CHE}"},
        "No obtuvo CHE": {"start": "{INICIO_RREE_NO_OBTUVO_NO_SOLICITO_CHE}", "end": "{FIN_RREE_NO_OBTUVO_NO_SOLICITO_CHE}"},
        "CHE en curso": {"start": "{INICIO_RREE_CHE_EN_CURSO}", "end": "{FIN_RREE_CHE_EN_CURSO}"
    },
    "RESIDUOS_GESTION": {
        "Correcta": {"start": "{INICIO_CORRECTA_GESTION_RESIDUOS}", "end": "{FIN_CORRECTA_GESTION_RESIDUOS}"
    },
    "ASP_STATUS": {
        "Finalizada": {"start": "{INICIO_PRESENTACION_ASP_FINALIZADA}", "end": "{FIN_PRESENTACION_ASP_FINALIZADA}"
    },
    "VALVULAS_CALIBRACION_STATUS": {
        "Cumple": {"start": "{INICIO_CALIBRACION_ASP_CUMPLE}", "end": "{FIN_CALIBRACION_ASP_CUMPLE}"
    },
    "ADA_STATUS": {
         "Prefactibilidad con Chi 0 para todos los permisos": {"start": "{PREFA_TODOS_CHI0}", "end": "{FIN_PREFA_TODOS_CHI0}"
         "Prefactibilidad vigente": {"start": "{PREFA_OBTENIDA}", "end": "{FIN_PREFA_OBTENIDA}"
         "Prefactibilidad vencida": {"start": "{PREFA_VENCIDA}", "end": "{FIN_PREFA_VENCIDA}"
         "No solicitada": {"start": "{PREFA_NO_SOLICITADA}", "end": "{FIN_PREFA_NO_SOLICITADA}"
    },
    "RENPRE_STATUS": {
        "No aplica": {"start": "{INICIO_NO_APLICA}", "end": "{FIN_NO_APLICA}"
        "Esta inscripto y renueva": {"start": "{INICIO_APLICA_INSCRIPTO_RENUEVA}", "end": "{FIN_APLICA_INSCRIPTO_RENUEVA}"
        "Esta inscripto y no renovó": {"start": "{INICIO_APLICA_INSCRIPTO_NO_RENOVO}", "end": "{FIN_APLICA_NO_RENOVO}"
        "Aplica pero no esta inscripto": {"start": "{INICIO_APLICA_NO_INSCRIPTO}", "end": "{FIN_APLICA_NO_INSCRIPTO}"
    },
    "SEGURO_STATUS": {
        "Vigente": {"start": "{INICIO_POLIZA_VIGENTE}", "end": "{FIN_POLIZA_VIGENTE}"
        "Vencida": {"start": "{INICIO_POLIZA_VENCIDA}", "end": "{FIN_POLIZA_VENCIDA}"
        "Nunca Tuvo": {"start": "{INICIO_NUNCA_TUVO_POLIZA}", "end": "{FIN_NUNCA_TUVO_POLIZA}"
    },
    "ACUMAR_STATUS": {
        "Vigente": {"start": "{INICIO_ACUMAR_VIGENTE}", "end": "{FIN_ACUMAR_VIGENTE}"
        "En curso": {"start": "{INICIO_ACUMAR_EN_CURSO}", "end": "{FIN_ACUMAR_EN_CURSO}"
        "Vencida": {"start": "{INICIO_ACUMAR_VENCIDA}", "end": "{FIN_ACUMAR_VENCIDA}"
        "No solicitada": {"start": "{INICIO_ACUMAR_NO_SOLICITADA}", "end": "{FIN_ACUMAR_NO_SOLICITADA}"
        "No aplica": {"start": "{INICIO_ACUMAR_NO_APLICA}", "end": "{FIN_ACUMAR_NO_APLICA}"
    },
    "INSCRIPCION_1102": {
        "Inscripto": {"start": "{INICIO_INSCRIPTA_1102}", "end": "{FIN_INSCRIPTA_1102}"
        "No inscripto": {"start": "{INICIO_NOINSCRIPTA_1102}", "end": "{FIN_NOINSCRIPTA_1102}"
        "No aplica": {"start": "{INICIO_NOAPLICA_1102}", "end": "{FIN_NOAPLICA_1102}"
    },
    "AUDITORIA_404": {
        "Realizo": {"start": "{INICIO_REALIZO_AUDITORIA}", "end": "{FIN_REALIZO_AUDITORIA}"
        "No realizo": {"start": "{INICIO_NO_REALIZO_AUDITORIA}", "end": "{FIN_NO_REALIZO_AUDITORIA}"
        "No aplica": {"start": "{INICIO_NOAPLICA_AUDITORIA}", "end": "{FIN_NOAPICA_AUDITORIA}"
        "No inscripto, no realiza": {"start": "{INICIO_NO_INSCRIPTA_NO_AUDITORIA}", "end": "{FIN_NO_INSCRIPTA_NO_AUDITORIA}"
    },
    "INSCRIPCION_277": {
        "Inscripto": {"start": "{INICIO_INSCRIPTA_277}", "end": "{FIN_INSCRIPTA_277}"
        "No inscripto": {"start": "{INICIO_NOINSCRIPTA_277}", "end": "{FIN_NOINSCRIPTA_277}"
        "No aplica": {"start": "{INICIO_NOAPLICA_277}", "end": "{FIN_NOAPLICA_277}"
    }
}

# ---------------------------
# UTILIDADES DE DOCX
# ---------------------------
def remove_paragraph(paragraph):
    """Elimina un párrafo del documento (python-docx)"""
    element = paragraph._element
    element.getparent().remove(element)


def find_paragraphs_to_remove(doc, selected_state, situation_type):
    """
    Reproduce la lógica que tenías en app.py:
    - Lee los marcadores start/end de MARCADORES_CONDICIONALES[situation_type]
    - Elimina las secciones que no correspondan al selected_state
    """
    paragraphs_to_remove = []
    markers_config = MARCADORES_CONDICIONALES.get(situation_type, {})

    selected_start_marker_text = markers_config.get(selected_state, {}).get('start')
    all_start_markers = {cfg.get('start') for cfg in markers_config.values() if cfg.get('start')}
    all_end_markers = {cfg.get('end') for cfg in markers_config.values() if cfg.get('end')}

    in_unselected_section = False

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            if in_unselected_section:
                paragraphs_to_remove.append(p)
            continue

        if text in all_start_markers:
            paragraphs_to_remove.append(p)
            if text != selected_start_marker_text:
                in_unselected_section = True
            else:
                in_unselected_section = False
            continue

        if text in all_end_markers:
            paragraphs_to_remove.append(p)
            in_unselected_section = False
            continue

        if in_unselected_section:
            paragraphs_to_remove.append(p)

    return paragraphs_to_remove


def find_and_remove(doc, key, section):
    """Helper: llama a find_paragraphs_to_remove y borra los párrafos encontrados"""
    try:
        for p in find_paragraphs_to_remove(doc, key, section):
            remove_paragraph(p)
    except Exception:
        # Si no hay configuración para esa sección, continuar
        pass


def reemplazar_marcadores(doc, user_data):
    """
    Reemplaza placeholders {KEY} en párrafos y tablas con los valores de user_data.
    Resalta en amarillo las variables que quedaron sin reemplazar (p.ej. {SOME_MARKER})
    """
    placeholder_regex = re.compile(r'{\s*([A-Z0-9_ÑÁÉÍÓÚ]+)\s*}', re.IGNORECASE)

    def process_container_paragraphs(paragraphs):
        for p in paragraphs:
            if not p.runs:
                continue
            combined_text = "".join([r.text for r in p.runs])
            modified_text = combined_text
            replaced_in_paragraph = False

            for marker, value in user_data.items():
                pattern = re.compile(r'{\s*' + re.escape(marker) + r'\s*}', re.IGNORECASE)
                if pattern.search(modified_text):
                    if str(value).strip() in ['N/A', '0', '']:
                        modified_text = pattern.sub('', modified_text)
                    else:
                        modified_text = pattern.sub(str(value), modified_text)
                    replaced_in_paragraph = True

            if replaced_in_paragraph:
                # eliminar runs originales
                for i in range(len(p.runs)-1, -1, -1):
                    try:
                        p.runs[i]._element.getparent().remove(p.runs[i]._element)
                    except Exception:
                        pass
                new_run = p.add_run(modified_text)

            # después de reemplazos, resaltar placeholders que siguen presentes
            for run in p.runs:
                if placeholder_regex.search(run.text):
                    try:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    except Exception:
                        pass

    for p in doc.paragraphs:
        process_container_paragraphs([p])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_container_paragraphs(cell.paragraphs)


def generar_tabla_desde_interfaz_dinamica(doc, lista_filas):
    """
    Inserta una tabla bajo el párrafo que contiene "Plan de monitoreos".
    lista_filas: lista de dicts con keys: 'recurso','organismo','puntos','parametros','frecuencia'
    """
    target_text = "Plan de monitoreos"
    encontrado = False

    for p in doc.paragraphs:
        if target_text in p.text:
            encontrado = True
            headers = ['Recurso', 'Organismo', 'Puntos', 'Parámetros', 'Frecuencia']
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, name in enumerate(headers):
                hdr_cells[i].text = name
                try:
                    run = hdr_cells[i].paragraphs[0].runs[0]
                    run.bold = True
                    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass

            for fila in lista_filas:
                recurso = fila.get('recurso', '').strip()
                if recurso:
                    row = table.add_row().cells
                    row[0].text = recurso
                    row[1].text = fila.get('organismo','')
                    row[2].text = fila.get('puntos','')
                    row[3].text = fila.get('parametros','')
                    frec = fila.get('frecuencia','')
                    row[4].text = frec if frec != 'Seleccione...' else ""
                    # set font size
                    for cell in row:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                try:
                                    run.font.size = Pt(10)
                                except Exception:
                                    pass
            break
    # si no encontrado, no hacer nada (el caller decide qué informar)


def insertar_tabla_manual_dinamica(doc, lista_de_filas_widgets):
    # Alias / compatibilidad (igual que generar_tabla_desde_interfaz_dinamica)
    generar_tabla_desde_interfaz_dinamica(doc, lista_de_filas_widgets)


def agregar_hallazgo_formateado_al_doc(doc, index, observacion, situacion, autoridad, riesgo, recomendacion):
    if not (str(observacion).strip() or str(situacion).strip()):
        return
    p_title = doc.add_paragraph()
    run_title = p_title.add_run(f"Observación de campo # {index}")
    run_title.bold = True
    run_title.underline = True
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    fields = [
        ("Observación:", observacion), ("Situación:", situacion), ("Autoridad:", autoridad),
        ("Riesgo:", riesgo), ("Recomendación:", recomendacion)
    ]

    doc.add_paragraph()
    for label, value in fields:
        if str(value).strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_label = p.add_run(label + " ")
            run_label.bold = True
            p.add_run(str(value))

# ---------------------------
# FUNCION PÚBLICA
# ---------------------------

def generate_report(template_bytes, user_data, hallazgos, muestreo_rows):
    """
    Genera el informe y devuelve un tuple (out_bytes, filename)
    - template_bytes: bytes de la plantilla .docx (desde upload)
    - user_data: dict con todos los campos a reemplazar (keys en mayúscula según plantilla)
    - hallazgos: lista de dicts con keys: observacion, situacion, autoridad, riesgo, recomendacion
    - muestreo_rows: lista de dicts con keys: recurso, organismo, puntos, parametros, frecuencia
    """
    doc = Document(io.BytesIO(template_bytes))

    # 1) Procesar logs condicionales (intentar remover bloques según los campos que están en user_data)
    # Llamamos a find_and_remove para secciones clave (si la clave no existe, la función absorbe excepción)
    try:
        find_and_remove(doc, user_data.get('HABILITACION_MUNICIPAL_STATUS', ''), "HABILITACION_MUNICIPAL")
    except Exception:
        pass
    try:
        find_and_remove(doc, user_data.get('CNCA_DETALLE_STATUS', ''), "CNCA_STATUS")
    except Exception:
        pass
    try:
        find_and_remove(doc, user_data.get('CAAP_STATUS', user_data.get('CAAP_LOGICA','')), "CAAP_STATUS")
    except Exception:
        pass
    try:
        # permitimos el mapeo directo del valor CAAF
        find_and_remove(doc, user_data.get('CAAF_STATUS', ''), "CAAF_STATUS")
    except Exception:
        pass
    # Otros
    mappings_to_check = [
        ('LEGA_STATUS', 'LEGA_STATUS'),
        ('RESIDUOS_ESPECIALES_STATUS','RESIDUOS_ESPECIALES_STATUS'),
        ('CHE_STATUS','CHE_STATUS'),
        ('GESTION_RESIDUOS_STATUS','RESIDUOS_GESTION'),
        ('ASP_STATUS','ASP_STATUS'),
        ('VALVULAS_STATUS','VALVULAS_CALIBRACION_STATUS'),
        ('ADA_STATUS','ADA_STATUS'),
        ('RENPRE_STATUS','RENPRE_STATUS'),
        ('SEGURO_STATUS','SEGURO_STATUS'),
        ('ACUMAR_STATUS','ACUMAR_STATUS'),
        ('SE_STATUS','SE_STATUS'),
        ('INSCRIPCION_1102','INSCRIPCION_1102'),
        ('AUDITORIA_404','AUDITORIA_404'),
        ('INSCRIPCION_277','INSCRIPCION_277')
    ]
    for key, section in mappings_to_check:
        try:
            find_and_remove(doc, user_data.get(key, ''), section)
        except Exception:
            pass

    # 2) Reemplazar marcadores con user_data
    reemplazar_marcadores(doc, user_data)

    # 3) Insertar hallazgos
    for i, h in enumerate(hallazgos, start=1):
        agregar_hallazgo_formateado_al_doc(
            doc, i,
            h.get('observacion',''), h.get('situacion',''),
            h.get('autoridad',''), h.get('riesgo',''), h.get('recomendacion','')
        )

    # 4) Insertar tabla de muestreo
    if muestreo_rows:
        generar_tabla_desde_interfaz_dinamica(doc, muestreo_rows)

    # 5) Guardar a bytes y devolver
    out_stream = io.BytesIO()
    razon = user_data.get('RAZON_SOCIAL', 'Empresa').replace(' ','_')
    filename = f"Informe_{razon}.docx"
    doc.save(out_stream)
    out_stream.seek(0)
    return out_stream.read(), filename
